import streamlit as st
import os
import re
import tempfile
from dotenv import load_dotenv
import logging
from typing import Dict, Any, Optional, Tuple, Union
from utils import fetch_models, ProviderHandler
from pydantic import Field, ValidationError
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
# LangChain imports
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader, TextLoader
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_groq import ChatGroq
from langchain_ollama import ChatOllama
from langchain.chains import LLMChain

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()


# Constants
MAX_FILE_SIZE_MB = 200  # 10MB maximum file size
SUPPORTED_FILE_TYPES = ["pdf", "docx", "txt", "md"]
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
DEFAULT_CHUNK_SIZE = 1000
DEFAULT_CHUNK_OVERLAP = 100

class Config:
    """Application configuration constants"""
    MARKETING_TASKS = [
        "Marketing Strategy",
        "Campaign Ideas",
        "Social Media Content",
        "SEO Optimization",
        "Copywriting"
    ]
    
    API_KEYS = {
        "GROQ": os.getenv("GROQ_API_KEY", ""),
        "OPENAI": os.getenv("OPENAI_API_KEY", ""),
        "PANDASAI": os.getenv("PANDAS_API_KEY", "")
    }

        
def get_api_key(provider: str) -> str:
    config = Config()
    return config.API_KEYS.get(provider.upper(), "")

# Initialize session state
def initialize_session_state():
    """Initialize Streamlit session state variables"""
    if "initialized" not in st.session_state:
        st.session_state.update({
            "initialized": True,
            "vector_store": None,
            "llm": None,
            "doc_content": "",
            "brand_description": "",
            "target_audience": "",
            "products_services": "",
            "marketing_goals": "",
            "existing_content": "",
            "keywords": "",
            "suggested_topics": "",
            "error_message": "",
            "processing_done": False,
            "task": "Marketing Strategy"  # Initialize task
        })

def convert_to_docx(content: str) -> bytes:
    """Convert markdown content to DOCX format"""
    document = Document()
    content = content.strip()
    
    # Split content into paragraphs
    paragraphs = content.split('\n\n')
    
    for paragraph in paragraphs:
        if paragraph.startswith('# '):
            # Heading
            heading = document.add_heading(paragraph[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif paragraph.startswith('## '):
            # Subheading
            heading = document.add_heading(paragraph[3:], level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif paragraph.startswith('- '):
            # Bullet point
            document.add_paragraph().add_run(paragraph).font.size = Pt(10)
        else:
            # Regular paragraph
            document.add_paragraph(paragraph)
            
    # Save to bytes
    file_bytes = BytesIO()
    document.save(file_bytes)
    file_bytes.seek(0)
    return file_bytes.read()


def create_sidebar() -> Dict[str, Any]:
    """Create the sidebar UI components"""
    default_endpoint = {
                "GROQ": "https://api.groq.com/openai/v1",
                "OPENAI": "https://api.openai.com/v1",
                "OLLAMA": "http://localhost:11434"
            }

    model = None

    with st.sidebar:

        st.header("🎯 Marketing Task")
        task = st.selectbox(
            "Select Task", 
            options=Config.MARKETING_TASKS,
            index=0
        )

        st.session_state.task = task

        st.title("⚙️ AI Configuration")
        
        provider = st.selectbox(
            "AI Provider", 
            options=["Groq", "Ollama"],
            key="provider_select",
            help="Select your preferred AI service provider"
        )
        
        with st.expander("Provider Settings", expanded=True):
         
            # API Configuration
            if provider == "Groq":
                default_endpoint = "https://api.groq.com/openai/v1"
            else:  # Ollama
                default_endpoint = "http://localhost:11434"


            endpoint = st.text_input(
                "API Endpoint",
                value=default_endpoint,
                key="endpoint_input"
            )
            st.session_state.endpoint = endpoint

            api_key = None
            

            if provider != "Ollama":
                api_key = st.text_input(
                    f"{provider} API Key",
                    type="password",
                    value=get_api_key(provider),
                    help=f"Get your API key from {provider}'s dashboard"
                )
                st.session_state.api_key = api_key
                st.sidebar.markdown("[Get Groq API Key](https://console.groq.com/keys)")
            else:
                st.session_state.api_key = None
                
                st.sidebar.markdown("[Download Ollama](https://ollama.com/)")
                
                           

            # Model selection with caching
        if provider == "Ollama" or st.session_state.get("api_key"):
            with st.spinner("Loading models..."):
                models = fetch_models(
                    provider,
                    st.session_state.get("endpoint"),
                    st.session_state.get("api_key")
                )
               
            model = st.selectbox(
                "Select AI Model",
                models,
                key="model_select",
                help="Select the model version to use"
            )
            st.session_state.model = model
        
        with st.expander("Advanced Settings", expanded=False):
            temperature = st.slider(
                "Temperature", 0.0, 1.0, 0.3,
                help="high temp = high creativity"
            )
            max_tokens = st.number_input(
                "Max Tokens", 
                min_value=512, 
                max_value=4096, 
                value=2048
            )
                        
        st.sidebar.markdown(
                            """<div style="text-align: center; margin-top: 20px;">
                                <a href="https://buymeacoffee.com/oguso">
                                    <img src="https://cdn.buymeacoffee.com/buttons/v2/default-yellow.png" alt="Buy Me A Coffee" style="width: 150px; height: auto;">
                                </a>
                                <p style="color: #666; margin-top: 5px;">Support my work!</p>
                            </div>
                            """, unsafe_allow_html=True
                            )
        
    return {
        "provider": provider,
        "model": model,
        "api_key": api_key,
        "api_endpoint": endpoint,
        "temperature": temperature,
        "max_tokens": max_tokens,
        "task": task,
    }

@st.cache_resource(show_spinner=False)
def initialize_llm(config: Dict[str, Any]) -> Optional[Union[ChatGroq, ChatOllama]]:
    """Initialize the language model with caching"""
    try:
        if config["provider"] == "Groq":
            if not config["api_key"]:
                st.error("Groq API key is required")
                st.info("[Get Groq API Key](https://console.groq.com/keys)")
                return None

            return ChatGroq(
                api_key=config["api_key"],
                model_name=config["model"],
                temperature=config["temperature"],
                max_tokens=config["max_tokens"]
            )
        else:
            return ChatOllama(
                model=config["model"],
                base_url=config["api_endpoint"],
                temperature=config["temperature"],
                num_predict=config["max_tokens"]
            )
    except Exception as e:
        logger.error(f"LLM initialization error: {str(e)}")
        st.error(f"Failed to initialize model: {str(e)}")
        return None

def validate_uploaded_file(file: st.runtime.uploaded_file_manager.UploadedFile) -> bool:
    """Validate uploaded file size and type"""
    if file.size > MAX_FILE_SIZE_MB * 1024 * 1024:
        st.error(f"File size exceeds {MAX_FILE_SIZE_MB}MB limit")
        return False
    
    file_extension = file.name.split('.')[-1].lower()
    if file_extension not in SUPPORTED_FILE_TYPES:
        st.error(f"Unsupported file type: {file_extension}")
        return False
    
    return True

@st.cache_data(show_spinner="Processing document...")
def process_document(_file: bytes, file_name: str) -> Tuple[Optional[FAISS], str]:
    """Process uploaded document and create vector store"""
    try:
        # Read the file content directly from the bytes
        file_extension = file_name.split('.')[-1].lower()

        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as temp_file:
            temp_file.write(_file)
            temp_path = temp_file.name  # Get the file path
        
        if file_extension == 'pdf':
            loader = PyPDFLoader(temp_path)
        elif file_extension in ['docx', 'doc']:
            loader = Docx2txtLoader(temp_path)
        elif file_extension in ['txt', 'md']:
            loader = TextLoader(temp_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
            
        documents = loader.load()
        doc_content = " ".join([doc.page_content for doc in documents])

        # Text splitting
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=DEFAULT_CHUNK_SIZE,
            chunk_overlap=DEFAULT_CHUNK_OVERLAP
        )
        splits = text_splitter.split_documents(documents)

        # Create embeddings
        embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)
        vector_store = FAISS.from_documents(splits, embeddings)

        return vector_store, doc_content

    except Exception as e:
        logger.error(f"Document processing failed: {str(e)}")
        st.error(f"Document processing error: {str(e)}")
        return None, ""


def generate_insights(llm: Any, vector_store: FAISS, field_name: str) -> str:
    """Generate all marketing insights using RAG"""
    retriever = vector_store.as_retriever(search_kwargs={"k": 3})
    
    # Create prompts for different fields
    field_prompts = {
        "brand_description": "Based on the provided context, write a concise brand description. Extract information about the company's mission, values, and unique selling points.",
        "target_audience": "Based on the provided context, identify and describe the target audience or customer segments for this business. Include demographics, psychographics, and key characteristics.",
        "products_services": "Based on the provided context, list and briefly describe the main products and/or services offered by the business.",
        "marketing_goals": "Based on the provided context, identify the key marketing goals or objectives for this business. If not explicitly stated, suggest reasonable goals based on the business type and information provided.",
        "existing_content": "Based on the provided context, summarize any existing marketing content, campaigns, or channels mentioned in the document.",
        "keywords": "Based on the provided context, generate a list of 10-15 relevant keywords for this business that could be used for marketing purposes. Format as a comma-separated list.",
        "suggested_topics": "Based on the provided context, suggest 5-7 content topics that would be relevant for this business's marketing strategy. Present as a numbered list."
    }

    prompt_template = """
    You are a marketing specialist tasked with analyzing business documents.
        
        {input}
        
        Context:
        {context}
        
        Provide a clear, concise response focusing only on the information requested.
    """
    
    document_chain = create_stuff_documents_chain(
        llm,
        ChatPromptTemplate.from_template(prompt_template),
        document_variable_name="context"
    )
    
    retrieval_chain = create_retrieval_chain(retriever, document_chain)
    
    try:
        # Execute the chain
        result = retrieval_chain.invoke({
            "input": field_prompts[field_name]
        })
        return parse_insights(field_name, result["answer"])
    except Exception as e:
        logger.error(f"Insight generation failed: {str(e)}")
        return ""

def parse_insights(field_name: str, text: str) -> str:
    """
    Parses the LLM-generated text for a single field.

    - Trims unnecessary whitespace.
    - Ensures consistent formatting.
    - Handles lists, keywords, and bullet points correctly.

    :param field_name: The name of the field being processed.
    :param text: The raw LLM response.
    :return: Cleaned and formatted text.
    """
    text = text.strip()  # Remove leading/trailing spaces

    # Special handling for keyword-based fields (convert to comma-separated)
    if field_name.lower() in {"keywords", "suggested_topics"}:
        lines = [line.strip("-• ") for line in text.splitlines() if line.strip()]
        return ", ".join(lines)

    return text  # Return cleaned text for normal fields

def create_marketing_form() -> Dict[str, str]:
    """Create and handle the main marketing input form"""
    with st.form("marketing_form"):
        st.header("Business Information")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Brand Identity")
            brand = st.text_area(
                "Description",
                key="brand_description_input",
                value=st.session_state.get("brand_description", ""),
                height=150
            )
            
            st.subheader("Target Audience")
            audience = st.text_area(
                "Target Audience",
                key="target_audience_input",
                value=st.session_state.get("target_audience", ""),
                height=150
            )

            st.subheader("Existing Content")
            existing_content = st.text_area(
                "Existing Content",
                key="existing_content_input",
                value=st.session_state.get("existing_content", ""),
                height=150
            )

                    
        with col2:
            st.subheader("Products/Services")
            products = st.text_area(
                "Products/Services",
                key="products_services_input",
                value=st.session_state.get("products_services", ""),
                height=150
            )
            
            st.subheader("Marketing Goals")
            goals = st.text_area(
                "Marketing Goals",
                key="marketing_goals_input",
                value=st.session_state.get("marketing_goals", ""),
                height=150
            )

            st.subheader("SEO Content Strategy")
            keywords = st.text_area(
                "Keywords",
                key="keywords_input",
                value=st.session_state.get("keywords", ""),
                height=150
            )

        

        st.subheader("Media Communication")
        # Convert suggested_topics to a list of options
        suggested_topics = st.session_state.get("suggested_topics", "")
        topics_list = re.split(r'\d+\.\s*', suggested_topics.strip())

        topics_list = [topic.strip() for topic in topics_list if topic.strip()] if suggested_topics else []
        
        selected_topic = st.selectbox(
            "Suggested Topics",
            key="suggested_topics_input",
            options=["Select a topic"] + topics_list,
            help="Select a suggested topic for your content"
        )
        
        post_type = st.selectbox(
                "Post Type",
                options=["Instagram", "LinkedIn", "Twitter", "Blog", "Podcast", "Media Brief"],
                key="post_type_input",
                help="Select the type of post"
            )
        
        tone = st.selectbox(
                "Tone of Voice",
                options=["Formal", "Casual", "Professional", "Friendly"],
                key="tone_input",
                help="Select the tone for the generated content"
            )
        
        # Add a submit button
        if st.form_submit_button(f"🚀 Generate {st.session_state.task}"):
            return {
                "brand_description": brand,
                "target_audience": audience,
                "products_services": products,
                "marketing_goals": goals,
                "keywords": keywords,
                "suggested_topics": selected_topic,
                "existing_content": existing_content,
                "tone": tone,
                "post_type": post_type
            }
    return {}


def generate_output(llm: Any, task: str, form_data: Dict[str, str]) -> str:
    """Generate task-specific marketing content"""
    task_prompts = {
    "Marketing Strategy": """
        You are a senior marketing strategist tasked with creating a comprehensive marketing plan.
        
        ## Business Context
        Brand Description: {brand_description}
        Target Audience: {target_audience}
        Products/Services: {products_services}
        Marketing Goals: {marketing_goals}
        Existing Content: {existing_content}
        Keywords: {keywords}
        
        ## Instructions
        Develop a detailed, actionable marketing strategy that aligns with the business goals.
        Focus on creating a strategy that is specific, measurable, achievable, relevant, and time-bound.
        
        ## Required Output Structure
        1. Executive Summary (brief overview of the entire strategy)
        2. Market Analysis (industry trends, competitive landscape)
        3. Target Audience Segmentation (detailed profiles of key segments)
        4. Value Proposition & Positioning (unique selling points, brand positioning)
        5. Marketing Channels & Tactics (prioritized by ROI potential)
        6. Content Strategy (topics, formats, distribution, calendar)
        7. Budget Allocation (recommended spending by channel)
        8. Implementation Timeline (30-60-90 day plan)
        9. KPIs & Success Metrics (specific measurements for each goal)
        10. Risk Assessment & Contingency Plans
    """,
    
    "Campaign Ideas": """
        You are a creative campaign director tasked with developing innovative marketing campaigns.
        
        ## Business Context
        Brand Description: {brand_description}
        Target Audience: {target_audience}
        Products/Services: {products_services}
        Marketing Goals: {marketing_goals}
        Keywords: {keywords}
        Selected Topics: {suggested_topics}
        Tone: {tone}
        
        ## Instructions
        Generate 5 distinct, creative campaign concepts that align with the brand identity and will resonate with the target audience.
        Each campaign should be achievable with realistic resources and have clear business impact.
        
        ## Required Output Structure
        For each of the 5 campaigns, provide:
        
        ### Campaign [Number]: [Creative Name]
        * Concept: Brief explanation of the campaign idea and creative angle
        * Target Segment: Specific audience segment this will appeal to most
        * Core Message: The primary takeaway for the audience
        * Campaign Elements: List of deliverables (videos, posts, emails, etc.)
        * Channels: Primary platforms for distribution
        * Timeline: Suggested duration and key milestones
        * Success Metrics: How to measure campaign effectiveness
        * Estimated Impact: Expected outcomes tied to marketing goals
    """,
    
    "Social Media Content": """
        You are an expert social media manager creating platform-specific content.
        
        ## Business Context
        Brand Description: {brand_description}
        Target Audience: {target_audience}
        Products/Services: {products_services}
        Marketing Goals: {marketing_goals}
        Keywords: {keywords}
        Selected Topics: {suggested_topics}
        Tone: {tone}
        Post Type: {post_type}
        
        ## Instructions
        Create a comprehensive social media content plan optimized for {post_type}.
        Focus on engaging the target audience with content that drives specific marketing goals.
        Ensure all content maintains the brand's {tone} tone of voice.
        
        ## Required Output Structure
        1. Platform Strategy
           * Why {post_type} is effective for this audience
           * Best practices specific to this platform
           * Posting frequency recommendations
        
        2. Content Pillars (3-4 key themes aligned with business goals)
        
        3. Content Calendar (2-week sample)
           * Week 1:
             * Day 1: [Content type] - [Example post with exact copy]
             * Day 2: [Content type] - [Example post with exact copy]
             [Continue for all week]
           * Week 2: [Same format]
        
        4. Engagement Strategy
           * Response templates for common interactions
           * Community-building tactics
           * User-generated content opportunities
        
        5. Growth Tactics
           * Hashtag strategy (10-15 targeted hashtags grouped by purpose)
           * Collaboration opportunities
           * Cross-promotion ideas
        
        6. Analytics Focus
           * Key metrics to track for this specific platform
           * Benchmarks for success
    """,
    
    "SEO Optimization": """
        You are an SEO specialist developing a comprehensive search optimization strategy.
        
        ## Business Context
        Brand Description: {brand_description}
        Target Audience: {target_audience}
        Products/Services: {products_services}
        Marketing Goals: {marketing_goals}
        Keywords: {keywords}
        Existing Content: {existing_content}
        
        ## Instructions
        Create a detailed SEO strategy that will improve organic visibility and drive qualified traffic.
        Focus on both quick wins and long-term sustainable growth.
        Provide specific, actionable recommendations rather than general advice.
        
        ## Required Output Structure
        1. Keyword Strategy
           * Primary Keywords (5-7 high-priority terms with search volume estimates)
           * Secondary Keywords (10-15 supporting terms)
           * Long-tail Opportunities (7-10 specific phrases)
           * Semantic/Topic Clusters (group related terms by topic)
        
        2. On-Page Optimization
           * Title Tag Templates
           * Meta Description Frameworks
           * Heading Structure Recommendations
           * Content Length and Formatting Guidelines
           * Internal Linking Strategy
        
        3. Technical SEO Checklist
           * Site Speed Optimization
           * Mobile Usability
           * Schema Markup Recommendations
           * Indexation Controls
           * URL Structure Guidelines
        
        4. Content Strategy
           * Content Gaps Analysis
           * Content Update Priorities
           * New Content Recommendations (5-7 specific pieces)
           * Content Calendar Framework
        
        5. Off-Page Strategy
           * Link Building Tactics (specific to industry)
           * Digital PR Opportunities
           * Local Citation Opportunities (if applicable)
        
        6. Measurement Plan
           * Key Performance Indicators
           * Tracking Setup Recommendations
           * Reporting Schedule and Format
        """,

    "Copywriting": """
        You are a professional copywriter creating compelling {post_type} content.
        
        ## Business Context
        Brand Description: {brand_description}
        Target Audience: {target_audience}
        Products/Services: {products_services}
        Marketing Goals: {marketing_goals}
        Keywords: {keywords}
        Existing Content: {existing_content}
        Selected Topics: {suggested_topics}
        
        ## Instructions
        Create high-converting {post_type} copy that speaks directly to the target audience.
        Maintain a {tone} tone throughout while incorporating strategic keywords {keywords} naturally.
        The copy should directly support the stated marketing goals and the selected topics {suggested_topics}.
        
        ## Platform-Specific Guidelines
        ### Instagram:
        - a visual element description
        - Create a captivating caption (max 150 words)
        - Include a strong hook in the first line
        - Use 5-10 relevant hashtags
        - Add a clear call-to-action
        
        ### LinkedIn:
        - Professional but engaging tone
        - 3-5 short paragraphs with white space
        - Include industry insights or data points
        - End with a thoughtful question or clear CTA
        
        ### Twitter:
        - Concise messaging under 280 characters
        - Include relevant hashtags (2-3)
        - Consider a visual element description
        - Create a compelling reason to click/engage
        
        ### Blog:
        - Compelling headline with primary keyword
        - 800-1200 words with clear structure
        - H2 and H3 subheadings containing keywords
        - Introduction with hook and thesis
        - Body with valuable insights/examples
        - Conclusion with next steps or CTA
        
        ### Podcast:
        - Episode title and description
        - Show notes with timestamps
        - Key talking points and questions
        - Call-to-action for listeners
        
        ### Media Brief:
        - Headline and subheadline
        - Key message points (3-5 bulletpoints)
        - Supporting facts/statistics
        - Quote from company representative
        - Call-to-action and contact information
        
        ## Required Output Structure
        - Headline/Title: Attention-grabbing, keyword-rich
        - Main Content: Formatted appropriately for {post_type}
        - Call-to-Action: Clear next step for the audience
        - [For social posts] Hashtags: Strategically selected for reach
        """
    }
    
    try:
        # Create the prompt
        prompt = ChatPromptTemplate.from_template(task_prompts[task])
        
        # Create the chain
        chain = prompt | llm | StrOutputParser()
        
        # Execute the chain
        response = chain.invoke(form_data)
        
        return response
    except Exception as e:
        logger.error(f"Content generation failed: {str(e)}")
        return f"Error generating content: {str(e)}"

def main():
    """Main application flow"""
    initialize_session_state()
    st.set_page_config(page_title="AI Marketing Assistant", layout="wide")
    
    # Sidebar configuration
    config = create_sidebar()
    st.session_state.llm = initialize_llm(config)
    
    # Main content area
    st.title(f"📋 {config['task']} Generator")
    
    # File upload handling
    uploaded_file = st.file_uploader(
        "Upload business document (PDF, DOCX, TXT)", 
        type=SUPPORTED_FILE_TYPES,
        help="add documents to help the AI understand you/your business"
    )
    
    if uploaded_file and validate_uploaded_file(uploaded_file):
        # Reset processing flag if a new file is uploaded
        if ("last_uploaded_file" not in st.session_state or 
            st.session_state.last_uploaded_file != uploaded_file.name):
            st.session_state.processing_done = False
            st.session_state.last_uploaded_file = uploaded_file.name

        # Optionally add a manual reset button for re-running extraction
        if st.button("Extract Data"):
            st.session_state.processing_done = False

        vector_store, content = process_document(
            uploaded_file.getvalue(),
            uploaded_file.name
        )

        st.session_state.vector_store = vector_store
        st.session_state.doc_content = content

        if st.session_state.llm and st.session_state.vector_store and not st.session_state.processing_done:
            # Generate content for the fields
            with st.spinner("Extracting data..."):
                # Always update the value (or conditionally, if you prefer)
                st.session_state.brand_description = generate_insights(
                    st.session_state.llm, st.session_state.vector_store, "brand_description"
                )
                st.session_state.target_audience = generate_insights(
                    st.session_state.llm, st.session_state.vector_store, "target_audience"
                )
                st.session_state.products_services = generate_insights(
                    st.session_state.llm, st.session_state.vector_store, "products_services"
                )
                st.session_state.marketing_goals = generate_insights(
                    st.session_state.llm, st.session_state.vector_store, "marketing_goals"
                )
                st.session_state.existing_content = generate_insights(
                    st.session_state.llm, st.session_state.vector_store, "existing_content"
                )
                st.session_state.keywords = generate_insights(
                    st.session_state.llm, st.session_state.vector_store, "keywords"
                )
                st.session_state.suggested_topics = generate_insights(
                    st.session_state.llm, st.session_state.vector_store, "suggested_topics"
                )

                # Mark processing as done so we don't re-run it on subsequent reruns
                st.session_state.processing_done = True

            st.rerun()  # Force a re-run so the form picks up the new values

    
    # Main form and generation
    form_data = create_marketing_form()
    
    if form_data and st.session_state.llm:
        with st.spinner("Generating content..."):
            result = generate_output(
                st.session_state.llm,
                config["task"],
                form_data
            )
            
            st.subheader("Generated Content")
            st.markdown(result)
            
            # Add download button with format option
            
            docx_file = convert_to_docx(result)
            st.download_button(
                label="Download Result",
                data=docx_file,
                file_name=f"{config['task'].replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )


if __name__ == "__main__":
    main()